"""
DOCX document parser using python-docx library.
"""

import io
import logging
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)


class DOCXParseError(Exception):
    """Exception raised when DOCX parsing fails."""
    pass


class DOCXParser:
    """Parser for DOCX documents using python-docx library."""
    
    def __init__(self):
        """Initialize the DOCX parser."""
        pass
    
    def parse(self, content: bytes) -> str:
        """
        Parse DOCX content and extract text.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            DOCXParseError: If parsing fails
        """
        try:
            # Create a BytesIO object from the content
            docx_stream = io.BytesIO(content)
            
            # Create Document object
            doc = Document(docx_stream)
            
            # Extract text from all paragraphs
            text_content = []
            
            logger.info(f"Parsing DOCX with {len(doc.paragraphs)} paragraphs")
            
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:  # Only add non-empty paragraphs
                    text_content.append(paragraph_text)
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise DOCXParseError("No text content could be extracted from DOCX")
            
            # Join all content with newlines
            full_text = "\n".join(text_content)
            
            logger.info(f"Successfully parsed DOCX: {len(full_text)} characters extracted")
            
            return full_text
            
        except DOCXParseError:
            raise
        except Exception as e:
            raise DOCXParseError(f"Failed to parse DOCX: {str(e)}")
    
    def can_parse(self, content_type: str) -> bool:
        """
        Check if this parser can handle the given content type.
        
        Args:
            content_type: MIME type of the content
            
        Returns:
            True if parser can handle this content type
        """
        return content_type.lower() in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroenabled.12',
            'application/msword',
            'application/x-msword'
        ]